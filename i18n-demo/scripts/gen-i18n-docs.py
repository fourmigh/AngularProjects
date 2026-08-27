# -*- coding: utf-8 -*-
"""Generate beginner-friendly i18n docs for the Angular i18n-demo:
  1) docs/i18n三方案详解.docx  —— explanations of the three schemes
  2) docs/i18n三方案对比.xlsx  —— comparison matrix + glossary + file map + extras
Run from the project root:  python scripts/gen-i18n-docs.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

ROOT = Path(__file__).resolve().parent.parent
DOC_DIR = ROOT / "docs"
DOCX_PATH = DOC_DIR / "i18n三方案详解.docx"
XLSX_PATH = DOC_DIR / "i18n三方案对比.xlsx"

ACCENT = RGBColor(0x1F, 0x6F, 0xB2)
HIGHLIGHT = RGBColor(0x0B, 0x53, 0x94)
GREY = RGBColor(0x88, 0x88, 0x88)
CODE_BG = "F2F2F2"
TABLE_HEAD_BG = "1F6FB2"


# --------------------------------------------------------------------------
#  Word helpers
# --------------------------------------------------------------------------
def set_ea_font(run, name: str = "微软雅黑"):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def set_mono(run):
    run.font.name = "Consolas"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in (qn("w:ascii"), qn("w:hAnsi")):
        rfonts.set(attr, "Consolas")


def shade_cell(cell, fill=TABLE_HEAD_BG):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): fill})
    tcpr.append(shd)


def add_para(doc, text, style=None, size=None, bold=False, italic=False,
             color=None, align=None, space_after=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    if size:
        r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    set_ea_font(r)
    return p


def add_title(doc):
    add_para(doc, "Angular i18n 三种方案详解", size=22, bold=True, color=ACCENT,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "（基于 i18n-demo · 原理、机制、对比、架构与文件格式）", size=12, color=GREY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    r.font.size = Pt(16 if level == 1 else 13)
    r.font.color.rgb = HIGHLIGHT if level == 1 else ACCENT
    set_ea_font(r)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r0 = p.add_run(bold_prefix)
        r0.bold = True
        set_ea_font(r0)
    r = p.add_run(text)
    set_ea_font(r)
    return p


def add_code(doc, code: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, CODE_BG)
    lines = code.rstrip("\n").split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        r = p.add_run(line if line else " ")
        set_mono(r)
        r.font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_ea_font(r)
        shade_cell(cell)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            for k, seg in enumerate(str(val).split("\n")):
                p = cell.paragraphs[0] if k == 0 else cell.add_paragraph()
                r = p.add_run(seg)
                set_ea_font(r)
                if j == 0:
                    r.bold = True
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# --------------------------------------------------------------------------
#  Content constants (shared by Word and Excel)
# --------------------------------------------------------------------------
CODE_RUNTIME = r"""// 源码标记（TS）：$localize + 自定义消息 id @@demo.title
readonly title = $localize`:@@demo.title:Angular Native i18n Demo`;

// 语言切换（official 页）：运行时把译文注册进官方 registry
switchLanguage(id) {
  loadTranslations(map);                  // { 'demo.title': 'Angular 原生 i18n 演示', ... }
  this.renderTick.update(n => n + 1);     // 触发界面按新语言重建
}
"""

CODE_COMPILE = r"""# 官方编译时方案：extract-i18n + --localize
# 1) 扫描所有 $localize/i18n 标记，生成源消息文件
ng extract-i18n                    # -> src/locale/messages.xlf（<source> 原文）

# 2) 为每种语言填译文 <target>（本项目用 make-xlf.mjs 从 translations.json 自动注入）
#    -> messages.zh.xlf / messages.de.xlf

# 3) 按语言各构建一次，构建期直接内联译文
ng build -c compile                # angular.json: "compile": { "localize": true }
                                   # 产物: dist/.../browser/{,zh,de}，base href 分别 /、/zh/、/de/
"""

CODE_CUSTOM = r"""// 自研方案：t()/label() 按"键"读取，键名由 split-i18n 自动生成类型（拼错编译报错）
this.i18n.t('demo.welcome', { USER: this.name });   // 组件里不用手写 $localize
this.i18n.label('demo.items');

// 底层其实还是官方 $localize + loadTranslations：
export const SOURCE_MESSAGES = {
  'demo.title': () => $localize`:@@demo.title:Angular Native i18n Demo`,
  // ...
};
"""

CODE_XLF = r"""<?xml version="1.0" encoding="UTF-8" ?>
<xliff version="2.0" xmlns="urn:oasis:names:tc:xliff:document:2.0" srcLang="en-US">
  <file id="ngi18n" original="ng.template">
    <unit id="demo.title">                 <!-- 消息 id，别改 -->
      <segment>
        <source>Angular Native i18n Demo</source>   <!-- 原文（自动生成） -->
        <target>Angular 原生 i18n 演示</target>      <!-- 译文（译者/CAT 填） -->
      </segment>
    </unit>
  </file>
</xliff>
"""

CODE_OFFICIAL_JSON = r"""// 官方 --format=json：按"消息 id"索引，每个语言一个文件
{
  "zh": {
    "demo.title": "Angular 原生 i18n 演示",
    "demo.items": "你的购物车中有 {$COUNT} 件商品"
  }
}
"""

CODE_MARK_TS = r"""// TS 里：$localize 标签模板 + 消息 id，参数用 :名称: 声明
readonly welcome = $localize`:@@demo.welcome:Hello, ${this.userName}:USER:!`;
readonly items   = $localize`:@@demo.items:You have ${this.itemCount}:COUNT: items in your cart`;
"""

CODE_MARK_TEMPLATE = r"""<!-- 模板里：i18n 属性 + 自定义 @@id（方便与 JSON/xlf 对应） -->
<h1 i18n="@@demo.title">Angular Native i18n Demo</h1>
<p i18n="@@demo.subtitle">Built with @angular/localize</p>

<!-- 模板属性也支持：i18n-属性名 -->
<img [attr.alt]="logoAlt" i18n-alt="@@demo.alt" />
"""

CODE_MARK_ICU = r"""<!-- ICU 复数：模板里 -->
<span i18n="@@demo.cart">{{ count, plural, =0 {empty} =1 {1 item} other {{{count}} items}}}</span>

<!-- 提取进 xlf 后的 <target>：<x id="INTERPOLATION"/> 占位符只能移动、不能删除 -->
<unit id="demo.cart">
  <segment>
    <source>{VAR_PLURAL, plural, =0 {empty} =1 {1 item} other {<x id="INTERPOLATION"/> items}}</source>
    <target>{VAR_PLURAL, plural, =0 {空} =1 {1 件} other {<x id="INTERPOLATION"/> 件商品}}</target>
  </segment>
</unit>
"""

CODE_SWITCH = r"""// 语言状态：唯一数据源 LocaleService.locale（signal）
// 页面 effect：locale 变化 -> 切换官方注册表 -> 触发重建
effect(() => {
  const locale = this.localeService.locale();
  if (this.i18n.ready()) this.i18n.switchLanguage(locale);   // loadTranslations + renderTick++
});

// 模板用 renderTick 做 track，强制按新语言重算
@for (bounce of [i18n.renderTick()]; track bounce) { ... }
"""

CODE_BUILDCHAIN = r"""npm run build:compile
  1. i18n:check     校验翻译键（46/46，三语言齐全）
  2. i18n:extract   生成 src/locale/messages.xlf（67 条源消息）
  3. i18n:make:xl   注入译文 -> messages.zh.xlf / messages.de.xlf（46 个 target）
  4. ng build -c compile   localize:true
                     产物：dist/i18n-demo/browser/{,zh,de}（各含 index.html）
"""

ANALOGIES = {
    "runtime": "电子菜单：点「切换语言」按钮，界面当场换成另一种语言，不用重新上菜（不用重新构建）。",
    "compile": "提前把菜单印成几本语言版本，发哪本读哪本：构建时就把每句话翻好定死，运行时不能再切。",
    "custom": "一份随时可以改的「翻译总表」：改一行文字，应用里立刻生效，还自带词典校验和类型保护。",
}

SNAPSHOT_HEADERS = ["维度", "官方运行时 ($localize)", "官方编译时 (extract-i18n)", "自研 t() (JSON+codegen)"]
SNAPSHOT_ROWS = [
    ["一句话", "运行时即时切换语言", "构建期定死语言、每语言一个站点", "JSON 总表 + 类型安全 + 可视化编辑器"],
    ["翻译时机", "运行时 loadTranslations()", "构建期就地替换译文", "运行时 loadTranslations()"],
    ["产物", "单一 bundle（可切换）", "每语言独立静态站点", "单一 bundle（可切换）"],
    ["切换", "按钮即时切换", "固定，不可切", "按钮/编辑器即时切换"],
    ["加语言成本", "改 JSON + 重新构建", "重新 build + 翻译 xlf", "改 JSON 即可"],
    ["适合", "需要运行时切换/灵活", "性能/SEO/固定语言站点", "团队要类型安全 + 可编辑文案"],
]

XLF_JSON_HEADERS = ["对比项", "xlf（XLIFF）", "JSON"]
XLF_JSON_ROWS = [
    ["格式", "标准 XML（XLIFF 2.0）", "自定义 JSON 键值"],
    ["结构", "<unit id=…><source>/<target>，带元数据与占位符", "{ 键: { 语言: 文本 } }"],
    ["生成", "ng extract-i18n 自动生成", "手写维护 / 脚本生成"],
    ["使用时机", "构建期 --localize 内联", "运行时 loadTranslations"],
    ["单/多语言", "每语言一个文件（messages.zh.xlf / de.xlf），含 source+target", "一份文件放全部语言"],
    ["CAT / 翻译工具", "主流平台都支持", "不支持 CAT 生态"],
    ["本项目里", "编译时方案的翻译源", "运行时方案与自研方案的文案总表"],
]

FORMAT_HEADERS = ["格式", "扩展名", "一句话", "CAT 兼容", "占位符 / ICU", "适用生态"]
FORMAT_ROWS = [
    ["ARB", ".arb", "Google 的 Flutter 翻译格式，键值对", "部分平台支持", "支持", "Flutter / 多端"],
    ["JSON", ".json", "官方 JSON：{ locale: { 消息id: 译文 } }", "不支持", "支持", "小项目 / 脚本友好"],
    ["XLIFF 1.2", ".xlf", "老牌标准：<trans-unit> / <source> / <target>", "主流 CAT 都支持", "支持", "广泛使用"],
    ["XLIFF 2.0", ".xlf", "新标准：<unit> / <segment>；本项目 extract 用它", "兼容面较窄", "支持", "新项目"],
    ["XMB / XTB", ".xmb / .xtb", "XML Message Bundle：源 .xmb，译文 .xtb", "较少", "受限（ICU 不完整）", "国际化大厂"],
]

XLF_TOOLS_HEADERS = ["工具", "类型", "官网", "费用参考(2026)", "XLIFF 支持", "一句话特点"]
XLF_TOOLS_ROWS = [
    ["Phrase（原 Memsource）", "云端 TMS", "phrase.com", "译者版 $27/月起；团队版 $525/月起", "1.2 + 2.0", "CAT 编辑器口碑好，可和 GitHub 等联动"],
    ["Crowdin", "云端 TMS", "crowdin.com", "有免费档 / Pro 约 $50/月起", "1.2 + 2.0", "开源项目免费，xlf 导入导出方便"],
    ["Lokalise", "云端 TMS", "lokalise.com", "Explorer $144/月起", "1.2 + 2.0", "面向开发团队，改词即推送"],
    ["Transifex", "云端 TMS", "transifex.com", "Starter 约 $135/月（年付）起", "1.2 + 2.0", "AI 翻译 + 质量评分"],
    ["Smartcat", "云端", "smartcat.com", "译者免费；企业版 $249/月起", "1.2 + 2.0", "自带译者接单市场，个人译者免费"],
    ["Weblate", "开源自托管 / 云", "weblate.org", "自托管免费；托管 €39/月", "1.2（2.0 支持有限）", "开源、与 Git 深度集成"],
    ["memoQ", "桌面 CAT", "memoq.com", "译者版约 €396/年", "1.2 + 2.0", "Windows 专业工具，翻译记忆强"],
    ["SDL Trados Studio", "桌面 CAT", "trados.com", "Freelance 约 $52/月（年付）起", "1.2 + 2.0", "行业老牌，客户常指定"],
    ["OmegaT", "桌面开源", "omegat.org", "免费", "1.2 为主", "免费开源跨平台，零预算起步"],
]

XLIFF_VER_HEADERS = ["对比项", "XLIFF 1.2", "XLIFF 2.0"]
XLIFF_VER_ROWS = [
    ["元素结构", "<trans-unit><source>/<target>", "<unit><segment><source>/<target>"],
    ["extract 命令", "ng extract-i18n --format=xlf", "ng extract-i18n --format=xlf2"],
    ["CAT 工具兼容", "主流工具都支持", "支持较少（多数仍偏 1.2）"],
    ["本项目", "未使用", "默认产出 xlf2"],
]

ANGULAR_JSON_HEADERS = ["配置块", "作用"]
ANGULAR_JSON_ROWS = [
    ["i18n.sourceLocale", "声明源语言（en-US）与默认子路径（''）"],
    ["i18n.locales", "声明有哪些语言，每语言的 translation 文件路径与 subPath（/zh、/de）"],
    ["architect.extract-i18n", "extract-i18n 命令的配置：基于哪个 buildTarget、输出到哪、什么格式（xlf2）"],
    ["architect.build.configurations.compile", "编译时构建配置：localize:true + fileReplacements"],
]

ARCH_HEADERS = ["对比项", "URL 驱动（:lang）", "内存驱动（localStorage）"]
ARCH_ROWS = [
    ["语言在 URL", "是：/zh、/en、/de（首页、compile-time）", "否：/official、/custom"],
    ["切换方式", "导航到新 URL（go() 把前缀换成目标语言）", "setLocale() 内存切换，URL 不变"],
    ["刷新 / 分享", "URL 即语言，可直接分享", "localStorage 记忆，刷新恢复"],
    ["SEO", "语言内容带在 URL，利于收录", "无语言 URL，运行时才翻译"],
    ["路由守卫", "canMatch 只放行 en/zh/de 的 :lang 段", "无需"],
    ["demo 页面", "home、compile-time", "official、custom"],
]

SCRIPTS_HEADERS = ["npm 脚本", "命令", "作用", "属于"]
SCRIPTS_ROWS = [
    ["start", "i18n:check + i18n:split + ng serve", "开发服务器（4200）", "全部"],
    ["i18n:check", "node check-i18n-keys.mjs", "校验翻译键齐全、都被使用", "全部"],
    ["i18n:split", "node split-i18n.mjs", "生成 i18n-keys.ts 类型 + source-messages.ts 注册表", "自研"],
    ["i18n:extract", "ng extract-i18n", "生成 src/locale/messages.xlf（源消息）", "编译时"],
    ["i18n:make:xl", "node make-xlf.mjs", "把 zh/de 注入 messages.zh.xlf / messages.de.xlf", "编译时"],
    ["build", "i18n:check + i18n:split + ng build", "运行时单站点构建（含 monaco）", "运行时 / 自研"],
    ["build:compile", "check + extract + make:xl + ng build -c compile", "编译时三站点构建", "编译时"],
    ["preview:compile", "http-server dist/i18n-demo/browser", "本地预览编译时站点（默认 /zh）", "编译时"],
]

FR_HEADERS = ["被替换文件", "替换为", "作用"]
FR_ROWS = [
    ["src/app/app.routes.ts", "app.routes.compile.ts", "编译站只路由 CompileTimeComponent"],
    ["src/app/features.ts", "features.compile.ts", "RUNTIME_PAGES=false：隐藏运行时导航与首页卡"],
]

CHECKLIST_HEADERS = ["步骤", "操作", "涉及文件"]
CHECKLIST_ROWS = [
    ["1", "translations.json 新增语言列（如 fr）", "translations.json"],
    ["2", "跑 i18n:split / i18n:check，补全所有键的 fr 文本", "i18n-keys.ts（生成）"],
    ["3", "跑 i18n:extract 更新源消息", "src/locale/messages.xlf"],
    ["4", "跑 i18n:make:xl 生成 messages.fr.xlf", "src/locale/messages.fr.xlf"],
    ["5", "angular.json i18n.locales 注册 fr + subPath", "angular.json"],
    ["6", "跑 build:compile，得到 browser/fr/", "dist"],
    ["7", "部署 fr 子目录 + CDN / 语言路由", "部署环境"],
]

FAQ_HEADERS = ["问题", "原因", "排查 / 解决"]
FAQ_ROWS = [
    ["构建产物里搜不到中文译文", "minifier 把非 ASCII 存成 \\uXXXX 转义（如 \\u7f16\\u8bd1）", "用 grep 搜 \\uXXXX，或解码查看；三语言 main 哈希不同即已内联"],
    ["i18n:check 报错", "翻译键缺语言 / 多余 / 未被使用", "按提示补全或删除 translations.json 里的键，重跑 check"],
    ["messages.zh.xlf 缺 <target>", "make:xl 未注入或语言列缺值", "先 i18n:split / check，再重跑 i18n:make:xl；缺 target 时 --localize 回退原文"],
    ["编译站壳显示中文", "壳标签用了运行时 t()，默认 zh", "改用静态 $localize（随站点内联），开发态走注册表"],
    ["如何确认 --localize 内联了", "怀疑某语言站点内容不对", "看 index.html 的 lang、各语言 main 哈希/大小不同、编译站无运行时 registry"],
]

MARK_HEADERS = ["场景", "写法", "说明"]
MARK_ROWS = [
    ["TS 静态文本", "readonly title = $localize`:@@demo.title:Angular Native i18n Demo`;", "会被 extract-i18n 提取；dev 走注册表、编译站内联"],
    ["TS 带参数", "readonly welcome = $localize`:@@demo.welcome:Hello, ${name}:USER:!`;", "占位符 :USER: 变成 xlf 里的 <x id=\"USER\"/>，翻译时保留"],
    ["模板元素", "<h1 i18n=\"@@demo.title\">Angular Native i18n Demo</h1>", "i18n 属性 + 自定义 @@id，与 JSON/xlf 对应"],
    ["模板属性", "<img [attr.alt]=\"…\" i18n-alt=\"@@demo.alt\" />", "i18n-属性名 形式"],
    ["ICU 复数", "<span i18n=\"@@demo.cart\">{{ count, plural, =0 {…} other {…}}}</span>", "xlf 里成 <x id=\"INTERPOLATION\"/> 占位，翻译时保留"],
    ["自研 t()", "this.i18n.t('demo.title') / label('demo.title')", "键由 TranslationKey 类型约束，拼错编译报错"],
]

# --------------------------------------------------------------------------
#  Glossary + file map (Excel)
# --------------------------------------------------------------------------
GLOSSARY_HEADERS = ["术语", "通俗解释"]
GLOSSARY_ROWS = [
    ["i18n", "国际化（internationalization）的缩写：让程序支持多种语言。18 表示首尾字母间有 18 个字母。"],
    ["locale", "「地区 / 语言」标识，如 zh（中文）、en（英文）、de（德文）。"],
    ["$localize", "Angular 官方提供的标记函数，把一段文字标记成待翻译消息；运行时或构建时都会解析它。"],
    ["@@id / 消息 id", "每条翻译消息的唯一编号，如 @@demo.title；取译文时按它查表，改动原文会改变 id。"],
    ["messages.xlf / xlf2", "官方翻译文件（XML，XLIFF 格式），每个 unit 对应一条消息；<source> 是原文，<target> 是译文。"],
    ["extract-i18n", "Angular CLI 命令，扫描所有消息标记并生成源翻译文件。"],
    ["--localize / localize:true", "构建选项：构建时把译文内联替换 $localize() 调用，每种语言产出一份静态包。"],
    ["loadTranslations / clearTranslations", "运行时把译文注册 / 清空进官方翻译注册表的 API，可动态切语言。"],
    ["官方 registry", "官方翻译注册表 {消息id: 译文}，渲染时查它拿到当前语言文本。"],
    ["<source> / <target>", "xlf 里「原文」（source）与「译文」（target）；译者只改 target。"],
    ["renderTick / 重建次数", "demo 用来强制组件重新计算翻译的计数器（每次切语言 +1）。"],
    ["codegen（代码生成）", "脚本把翻译表自动生成 TS 类型与注册表，编译期就能查出键名错误。"],
    ["TranslationKey", "由生成得到的「合法翻译键」类型；写错键名编译直接报错。"],
    ["ICU / 占位符", "翻译里的复数、数字、选择等表达式，如 {count, plural, ...}；翻译时占位符不能删只能移动。"],
    ["base href / subPath", "站点子路径（如 /zh/、/de/），浏览器据此加载对应语言站点的资源。"],
    ["tree-shake", "构建时去掉用不到的代码；编译时方案每个语言站点只留该语言文字，包更小。"],
    ["localStorage", "浏览器本地存储；demo 用它记住语言选择，刷新后恢复。"],
    ["Intl / 本地化日期数字", "浏览器按 locale 排版日期、数字、货币的 API（demo 页有展示）。"],
]

EXTRA_GLOSSARY = [
    ["SEO（搜索引擎优化）", "让网页内容更容易被搜索引擎抓取、搜到并排到前面。编译时方案译文在构建时就固化进页面文件，利于收录；运行时方案内容要浏览器运行后才生成，较难收录。"],
    ["CDN（内容分发网络）", "把网站静态文件缓存到各地服务器，用户就近下载、加载更快，方便分语言分发。"],
    ["bundle（打包产物）", "构建工具把代码合并压缩成一批 JS/CSS 文件。「单一 bundle」指只出一份可切换语言的包。"],
    ["JSON", "一种文本数据格式，用 { 键: 值 } 描述数据；本项目的文案总表就是 JSON。"],
    ["静态站点", "构建时提前生成、无需后端动态处理的网页文件；编译时方案每个语言一个静态目录。"],
    ["编译时 / 运行时", "编译时（build time）指构建软件那一步；运行时（runtime）指软件在浏览器里跑起来之后。"],
    ["npm / npm run", "Node 的包管理器与运行脚本的命令；build:compile、i18n:make:xl 等脚本由它触发。"],
    ["CAT 工具（计算机辅助翻译）", "帮译者高效翻译的软件（Phrase、Transifex、Lokalise 等），常配合 xlf 使用。"],
    ["fetch", "浏览器内置的发起网络请求的函数，demo 用它在运行时加载 translations.json。"],
    ["Monaco 编辑器", "微软开源在线代码编辑器（VS Code 的内核），自研页用它做可视化改文案面板。"],
    ["API", "程序之间的接口/调用入口，例如 loadTranslations() 就是一个 API。"],
    ["XLIFF", "xlf 文件背后的标准格式名称（看 messages.xlf 词条）。"],
]

EXTRA_GLOSSARY2 = [
    ["ARB", "Google 定义的翻译格式（.arb），常用于 Flutter；也是 extract-i18n 的输出格式之一。"],
    ["XMB / XTB", "XML Message Bundle 翻译格式：源文件 .xmb，译文文件 .xtb。"],
    ["XLIFF 1.2", "xlf 的旧版本标准（trans-unit）；CAT 工具兼容最好，可用 --format=xlf 输出。"],
    ["canMatch", "Angular 路由守卫：只有满足条件才匹配该路由；demo 用它只放行 en/zh/de 的 :lang 段。"],
    ["signal / effect", "Angular 响应式原语：signal 是可订阅的状态，effect 在依赖的 signal 变化时自动重跑。"],
    ["fileReplacements", "构建配置：构建时把某文件替换成另一个；demo 用它把运行时路由/功能换成编译时版本。"],
    ["URL 驱动 / 内存驱动", "两种语言状态方案：语言写在 URL（:lang），或只存在内存/localStorage（/official、/custom）。"],
    ["$languages", "translations.json 里的可选字段，声明启用哪些语言；编辑器里可勾选。"],
]

GLOSSARY_ROWS = GLOSSARY_ROWS + EXTRA_GLOSSARY + EXTRA_GLOSSARY2

WORD_GLOSSARY_HEADERS = ["术语", "大白话"]
WORD_GLOSSARY_ROWS = EXTRA_GLOSSARY + EXTRA_GLOSSARY2

FILE_MAP_HEADERS = ["文件 / 目录", "作用", "属于"]
FILE_MAP_ROWS = [
    ["src/app/custom-i18n/i18n/translations.json", "三方案共用的文案总表（键 → {语言: 文本}）", "全部"],
    ["src/app/custom-i18n/scripts/make-xlf.mjs", "把 translations.json 的 zh/de 注入 messages.zh.xlf / messages.de.xlf", "编译时"],
    ["src/app/custom-i18n/scripts/split-i18n.mjs", "生成 i18n-keys.ts 类型 + source-messages.ts 注册表", "自研"],
    ["src/app/custom-i18n/scripts/check-i18n-keys.mjs", "校验所有翻译键都被使用、三种语言齐全", "全部"],
    ["src/locale/messages.xlf / messages.zh.xlf / messages.de.xlf", "官方提取 / 译好的翻译文件", "编译时"],
    ["src/app/official/official.component.ts + official-content.component.{ts,html}", "官方运行时页面（/official）", "官方运行时"],
    ["src/app/compile-time/compile-time.component.{ts,html}", "编译时页面（/compile-time），展示真实 xlf", "官方编译时"],
    ["src/app/custom-i18n/pages/custom-page.component.{ts,html}", "自研页面（/custom）= 左 demo + 右编辑器", "自研"],
    ["src/app/custom-i18n/demo/demo.component.{ts,html}", "自研 demo，用 t()/label() 读文案", "自研"],
    ["src/app/custom-i18n/editor/editor.component.{ts,html}", "Monaco JSON 编辑器，可视化改文案", "自研"],
    ["src/app/custom-i18n/i18n.service.ts", "自研核心服务：t()/label()、applyEdited、下载/重置", "自研"],
    ["src/app/custom-i18n/i18n-keys.ts", "生成的文件：翻译键类型（TranslationKey）", "自研"],
    ["src/app/custom-i18n/source-messages.ts", "生成的文件：消息注册表（键 → $localize）", "自研"],
    ["src/app/locale.service.ts", "语言状态（signal）+ localStorage 持久化", "全部"],
    ["src/app/app.component.{ts,html}", "应用壳：导航、语言按钮、URL/内存切换 go()", "全部"],
    ["src/app/home/home.component.{ts,html}", "首页卡片（/:lang）", "全部"],
    ["src/app/features.ts / features.compile.ts", "RUNTIME_PAGES 开关（编译站关闭运行时页）", "全部"],
    ["src/app/app.config.ts", "路由配置入口（provideRouter(routes)）", "全部"],
    ["src/app/app.routes.ts / app.routes.compile.ts", "运行时路由（含 /official /custom）与编译时路由", "全部"],
    ["angular.json", "i18n.locales、compile 配置（localize:true）、extract-i18n 配置、fileReplacements", "全部"],
    ["package.json", "脚本入口：i18n:check / i18n:split / build:compile 等", "全部"],
]


# --------------------------------------------------------------------------
#  Word build
# --------------------------------------------------------------------------
def build_docx():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    add_title(doc)

    # ---- 1. 基础概念
    add_heading(doc, "一、基础概念", 1)
    add_para(doc, "i18n（internationalization，国际化）就是让同一个程序能显示多种语言。Angular 提供官方方案，也可以自研方案。本 demo 演示了三种：")
    add_bullet(doc, "官方运行时方案 —— 用 $localize + loadTranslations。", bold_prefix="① ")
    add_bullet(doc, "官方编译时方案 —— 用 extract-i18n + --localize。", bold_prefix="② ")
    add_bullet(doc, "自研方案 —— 用 t()/label() + JSON 总表 + 类型生成。", bold_prefix="③ ")

    add_heading(doc, "先理解两个关键词", 2)
    add_para(doc, "「消息标记」：源码里写一条带唯一编号的翻译条目，例如 $localize`:@@demo.title:Angular Native i18n Demo`。其中 @@demo.title 是这条消息的编号（消息 id），冒号后面是默认（原文）文案。")
    add_para(doc, "「翻译注册表（registry）」：程序内部一张键值表，形如 { 消息id: 译文 }。$localize 或 t() 取文案时，按 id 去这张表里查。")

    add_heading(doc, "三者的共同点", 2)
    add_bullet(doc, "都使用同一套「消息 id」标记（$localize:@@…）；")
    add_bullet(doc, "文案都来自同一份 translations.json 总表（make-xlf 脚本再把它转成编译时用的 xlf）；")
    add_bullet(doc, "占位符 / ICU（复数、日期、时间等）三套方案都支持。")

    # ---- 2. xlf 与 JSON
    add_heading(doc, "二、xlf 与 JSON 翻译文件", 1)
    add_heading(doc, "xlf（XLIFF）是什么", 2)
    add_para(doc, "xlf 是 Angular 官方 i18n 的翻译文件格式（XLIFF 2.0，XML）。由 ng extract-i18n 从源码的 $localize 标记自动生成：每个 unit 一条消息，<source> 是原文、<target> 是译文。")
    add_code(doc, CODE_XLF)
    add_heading(doc, "xlf 怎么编辑", 2)
    add_bullet(doc, "平时由翻译平台 / CAT 工具（Phrase、Transifex、Lokalise、Crowdin、Weblate、memoQ、Trados）打开编辑，只填 <target>；")
    add_bullet(doc, "手动编辑：找到对应 unit，改或补一个 <target>；不改 unit 的 id，占位符 <x …> 与 ICU 只能移动位置、不能删除；")
    add_bullet(doc, "注意：多数 CAT 工具对 XLIFF 1.2（ng extract-i18n --format=xlf）支持更好，本项目 extract 用的是 xlf2（XLIFF 2.0）；")
    add_bullet(doc, "本项目实际不用手写：改 translations.json → 跑 npm run i18n:make:xl（make-xlf.mjs 自动把 zh/de 注入 <target>）。")
    add_heading(doc, "xlf 编辑工具一览（官网与费用）", 2)
    add_para(doc, "云端 TMS 适合团队协作（浏览器里分工、审校、管理）；桌面 CAT 是译者在本地用的专业工具，功能更强但通常只有 Windows。")
    add_table(doc, XLF_TOOLS_HEADERS, XLF_TOOLS_ROWS)
    add_para(doc, "注：费用为官网公开起价（2026 年核实），随套餐、账单方式（月付/年付）变化，请以官网为准；本项目实际用 make-xlf.mjs 自动填 <target>，并未接入这些工具。",
             italic=True, color=GREY)
    add_heading(doc, "xlf 与 JSON 的区别", 2)
    add_table(doc, XLF_JSON_HEADERS, XLF_JSON_ROWS)
    add_para(doc, "补充两点：① Angular 官方也支持 --format=json 的翻译文件，结构是 { locale: { 消息id: 译文 } }，和本项目自研总表 { 键: { zh/en/de } } 不是同一种；② 本项目里运行时方案读 translations.json，编译时方案读 xlf，两者靠同一个 @@id 对应起来。")

    # ---- 3. 官方文件格式全景
    add_heading(doc, "三、官方翻译文件格式全景", 1)
    add_para(doc, "ng extract-i18n 可以输出 5 种格式（--format 指定）。")
    add_table(doc, FORMAT_HEADERS, FORMAT_ROWS)
    add_heading(doc, "XLIFF 1.2 与 2.0 的区别", 2)
    add_table(doc, XLIFF_VER_HEADERS, XLIFF_VER_ROWS)
    add_heading(doc, "官方的 JSON 与自研 JSON 不是同一种", 2)
    add_code(doc, CODE_OFFICIAL_JSON)
    add_para(doc, "官方 JSON 按「消息 id」索引、每语言一个文件，能直接喂给 --localize；本项目的自研 JSON 是一份总表按「业务键」放全部语言，运行时由 loadTranslations 消费。")
    add_heading(doc, "extract-i18n 与 angular.json 的关系", 2)
    add_para(doc, "extract-i18n 是命令（动作），angular.json 是配置文件（声明）。命令按配置执行：提取后的文件名/路径必须和 i18n.locales 里 translation 指向的文件对上，--localize 构建才会读到译文。")
    add_table(doc, ANGULAR_JSON_HEADERS, ANGULAR_JSON_ROWS)

    # ---- 4. 方案一
    add_heading(doc, "四、方案一：官方运行时方案（$localize + loadTranslations）", 1)
    add_heading(doc, "生活化类比", 2)
    add_para(doc, ANALOGIES["runtime"])
    add_heading(doc, "原理与机制", 2)
    add_bullet(doc, "在源码里写 $localize 标记 + 消息 id；", bold_prefix="1) 标记：")
    add_bullet(doc, "Angular 把它编译成一个 $localize() 函数调用；", bold_prefix="2) 构建：")
    add_bullet(doc, "应用启动时 fetch translations.json，再把当前语言的译文用 loadTranslations() 注册进官方注册表；", bold_prefix="3) 加载：")
    add_bullet(doc, "界面渲染时按「消息 id」从注册表取当前语言文本；", bold_prefix="4) 渲染：")
    add_bullet(doc, "清掉旧译文 + 注册新语言 + renderTick++ 触发组件重建，页面当场刷新。", bold_prefix="5) 切换：")
    add_code(doc, CODE_RUNTIME)
    add_heading(doc, "演示中的体现", 2)
    add_bullet(doc, "访问 /official 页面；")
    add_bullet(doc, "有语言切换按钮、当前语言代码、本地化日期/数字（Intl）、组件重建计数。")
    add_heading(doc, "优点 / 缺点", 2)
    add_bullet(doc, "单站点即可多语言、切换即时无刷新、部署简单、改词只需改 JSON。", bold_prefix="优点：")
    add_bullet(doc, "包体与运行时开销略高、新增语言要重新构建、内容到运行时才翻译故对 SEO 不友好。", bold_prefix="缺点：")

    # ---- 5. 方案二
    add_heading(doc, "五、方案二：官方编译时方案（extract-i18n + --localize）", 1)
    add_heading(doc, "生活化类比", 2)
    add_para(doc, ANALOGIES["compile"])
    add_heading(doc, "原理与机制", 2)
    add_bullet(doc, "扫描所有消息标记，生成源消息文件 messages.xlf（<source> 是原文）；", bold_prefix="1) extract-i18n：")
    add_bullet(doc, "把源文件复制成每语言一份，逐条在 <target> 填译文（本项目由 make-xlf.mjs 从 translations.json 自动注入）；", bold_prefix="2) 翻译：")
    add_bullet(doc, "按语言各构建一次，构建期把 $localize(...) 直接替换成目标语言文本；", bold_prefix="3) build --localize：")
    add_bullet(doc, "每种语言一个独立静态站点（browser/{,zh,de}，base href 分别 /、/zh/、/de/）。", bold_prefix="4) 产物：")
    add_code(doc, CODE_COMPILE)
    add_heading(doc, "演示中的体现", 2)
    add_bullet(doc, "访问 /compile-time 页面；")
    add_bullet(doc, "页面会展示「真实提取出来的 messages.xlf」和「翻译后的示意」，并说明三站点的构建结果。")
    add_heading(doc, "优点 / 缺点", 2)
    add_bullet(doc, "包最小、可 tree-shake、每种语言独立站点利于 SEO 与 CDN 分发、无运行时 registry 开销。", bold_prefix="优点：")
    add_bullet(doc, "新增语言要重新 build + 翻译 xlf 并部署新目录、不能运行时切语言、翻译文件数量多。", bold_prefix="缺点：")

    # ---- 6. 方案三
    add_heading(doc, "六、方案三：自研 t()/label() 方案（JSON 总表 + 类型化 codegen + 可视化编辑器）", 1)
    add_heading(doc, "生活化类比", 2)
    add_para(doc, ANALOGIES["custom"])
    add_heading(doc, "原理与机制", 2)
    add_bullet(doc, "translations.json 是唯一文案总表：{ \"demo.title\": { \"zh\": \"…\", \"en\": \"…\", \"de\": \"…\" } }；", bold_prefix="1) 总表：")
    add_bullet(doc, "split-i18n.mjs 自动生成两份文件：① i18n-keys.ts（把「翻译键」变成 TypeScript 类型，拼错键名编译报错）② source-messages.ts（注册表：键 -> $localize）；", bold_prefix="2) codegen：")
    add_bullet(doc, "组件直接 this.i18n.t('翻译键') / label('key') 读取，底层仍是官方 $localize + loadTranslations；", bold_prefix="3) 读取：")
    add_bullet(doc, "页面右侧 Monaco 编辑器可直接改合并 JSON → 校验 → Apply → 重新注册译文 → 即时生效；可勾选语言（$languages）、下载、重置。", bold_prefix="4) 编辑器：")
    add_code(doc, CODE_CUSTOM)
    add_heading(doc, "编辑器「即时生效」的底层原理", 3)
    add_bullet(doc, "编辑器预填 + 自动存草稿：打开 /custom 时 Monaco 已放好合并总表（i18n.service.ts 的 getMergedContent()）；每次按键 onDidChangeModelContent 就把文本写进 localStorage 当草稿，没点 Apply 也能续写、刷新不丢。")
    add_bullet(doc, "点 Apply 触发 applyEdited(text)（i18n.service.ts:88）：先 parseMerged() 校验 JSON 结构与语言是否合法（zh/en/de 是否写错），不合法提示错误，合法才继续。")
    add_bullet(doc, "重新注册译文：合法后 applyLocale()（i18n.service.ts:227）用 buildMap(当前语言) 拼出 { 消息id: 译文 }，先 clearTranslations() 清空官方 $localize 注册表，再 loadTranslations(map) 装入新译文。")
    add_bullet(doc, "强制重画：注册表一更新，applyLocale() 让 renderTick 计数器 +1（i18n.service.ts:233）；而每个页面模板都用 @for (bounce of [i18n.renderTick()]; track bounce) 把整块包起来，renderTick 一变 Angular 就销毁并重建该视图，模板重新去注册表取译文。")
    add_bullet(doc, "立刻看到新字：模板重新执行 i18n.t('键') / $localize，这次从新注册表取到刚改的译文，画面当场刷新——不刷新页面、不重新打包。结果（及草稿）都存 localStorage，Reset 清空回到 translations.json 主文件。")
    add_heading(doc, "演示中的体现", 2)
    add_bullet(doc, "访问 /custom 页面；")
    add_bullet(doc, "左侧是 demo（随语言即时变化），右侧是可视化 JSON 编辑器。")
    add_heading(doc, "优点 / 缺点", 2)
    add_bullet(doc, "文案集中在一份总表、键名类型安全（拼错编译报错）、可在线改词即时生效、团队友好且不放弃官方机制。", bold_prefix="优点：")
    add_bullet(doc, "多一层封装、浏览器里带 Monaco 编辑器偏重、依赖生成脚本链。", bold_prefix="缺点：")

    # ---- 7. 对比速览
    add_heading(doc, "七、三方案对比速览", 1)
    add_table(doc, SNAPSHOT_HEADERS, SNAPSHOT_ROWS)
    add_para(doc, "更完整的对比矩阵、术语词典、源码文件地图，请打开同目录下的 Excel：i18n三方案对比.xlsx。")

    # ---- 8. 标记细节示例
    add_heading(doc, "八、标记细节示例", 1)
    add_para(doc, "消息标记可以写在两处：TS 里用 $localize 标签模板，模板里用 i18n 属性。两者都会进 extract-i18n 的扫描，也都能被运行时注册表解析。")
    add_code(doc, CODE_MARK_TS)
    add_code(doc, CODE_MARK_TEMPLATE)
    add_heading(doc, "ICU（复数 / 选择）在 xlf 里怎么翻译", 2)
    add_code(doc, CODE_MARK_ICU)
    add_para(doc, "要点：翻译 ICU 时保留 {VAR_PLURAL, plural, …} 骨架和 <x …> 占位符，只翻译每个分支里的文字；占位符可以调整位置，但不能删。")

    # ---- 9. 应用架构与语言状态
    add_heading(doc, "九、应用架构与语言状态", 1)
    add_para(doc, "语言状态（当前用哪种语言）在 demo 里有两种驱动方式。")
    add_table(doc, ARCH_HEADERS, ARCH_ROWS)
    add_heading(doc, "响应式切换机制", 2)
    add_bullet(doc, "LocaleService.locale（signal）是唯一数据源；")
    add_bullet(doc, "页面 effect 订阅 locale：变化 → i18n.switchLanguage → loadTranslations + renderTick++；")
    add_bullet(doc, "模板用 @for (bounce of [i18n.renderTick()]) 做 track，强制组件按新语言重算。")
    add_code(doc, CODE_SWITCH)
    add_heading(doc, "为什么壳导航标签用静态 $localize", 2)
    add_para(doc, "壳（顶部导航/语言按钮）如果走运行时 t()，在 --localize 静态站点里会因运行时默认 zh 而显示中文，破坏「每站固定语言」。改成静态 $localize 后：开发态从注册表取（随语言切换），编译态被 --localize 内联成站点语言。")

    # ---- 10. 工作流、构建与瘦身
    add_heading(doc, "十、工作流、构建与瘦身", 1)
    add_heading(doc, "npm 脚本一览", 2)
    add_table(doc, SCRIPTS_HEADERS, SCRIPTS_ROWS)
    add_heading(doc, "build:compile 完整流程", 2)
    add_code(doc, CODE_BUILDCHAIN)
    add_heading(doc, "fileReplacements：编译站瘦身", 2)
    add_para(doc, "编译时站点只需编译页，通过构建配置把「运行时路由/功能」替换掉，去掉官方运行时页与自研页（以及它们拉进来的 monaco），main 仅约 243KB、无 monaco 懒加载 chunk。")
    add_table(doc, FR_HEADERS, FR_ROWS)

    # ---- 11. 新增语言与部署实操
    add_heading(doc, "十一、新增语言与部署实操", 1)
    add_heading(doc, "新增一种语言（以 fr 为例）", 2)
    add_table(doc, CHECKLIST_HEADERS, CHECKLIST_ROWS)
    add_heading(doc, "编译时三站点怎么部署", 2)
    add_bullet(doc, "base href 由 --localize 按 subPath 自动生成（/zh/、/de/），浏览器据此加载资源；")
    add_bullet(doc, "把 dist/i18n-demo/browser 整个目录交给任意静态服务器即可；本地预览用 npm run preview:compile；")
    add_bullet(doc, "生产可配合 nginx 语言协商、CDN 按子路径分发；每语言独立 URL 对 SEO 友好。")

    # ---- 12. 常见问题与验证技巧
    add_heading(doc, "十二、常见问题与验证技巧", 1)
    add_table(doc, FAQ_HEADERS, FAQ_ROWS)

    # ---- 13. 选型建议
    add_heading(doc, "十三、选型建议", 1)
    add_bullet(doc, "要性能 / SEO / 每种语言单独发布 → 选官方编译时（方案二）。")
    add_bullet(doc, "需要用户随时切语言、单站点部署 → 选官方运行时（方案一）。")
    add_bullet(doc, "团队想统一管理文案、要类型安全、还要在线改词 → 在不放弃官方 $localize 的前提下用自研封装（方案三）。")

    # ---- 14. 常见名词速查
    add_heading(doc, "十四、常见名词速查", 1)
    add_table(doc, WORD_GLOSSARY_HEADERS, WORD_GLOSSARY_ROWS)
    add_para(doc, "其余术语（$localize、@@id、loadTranslations、tree-shake 等）也收录在 Excel「名词小词典」里。")

    # ---- 15. 说明
    add_heading(doc, "十五、说明", 1)
    add_para(doc, "文中代码为示意；实际代码在 src/app 对应组件里。三方案共用同一套消息 id 与 translations.json 文案源，这正是 make-xlf.mjs 做桥接的原因。",
             italic=True, color=GREY)

    DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_PATH))
    print(f"[docx] saved: {DOCX_PATH}")


# --------------------------------------------------------------------------
#  Excel build
# --------------------------------------------------------------------------
thin = Side(style="thin", color="BFBFBF")
border_all = Border(left=thin, right=thin, top=thin, bottom=thin)
header_fill = PatternFill("solid", fgColor="1F6FB2")
alt_fill = PatternFill("solid", fgColor="EAF3FB")
header_font = Font(color="FFFFFF", bold=True, size=11)
wrap = Alignment(wrap_text=True, vertical="top")
header_align = Alignment(vertical="center", horizontal="center")

COMPARE_HEADERS = ["维度", "官方运行时 ($localize + loadTranslations)", "官方编译时 (extract-i18n + --localize)", "自研 t() (JSON + codegen)"]
COMPARE_ROWS = [
    ["一句话定位", "单一站点、运行时即时切换语言", "每种语言构建一个固定语言站点", "JSON 总表 + 类型安全 + 可视化编辑器"],
    ["通俗比喻", "电子菜单，现场切语言", "菜单提前印好多语言，发哪本读哪本", "可编辑的翻译总表，改完即时生效"],
    ["翻译何时生效", "运行时（浏览器里 loadTranslations()）", "构建期（--localize 把译文内联进 JS）", "运行时（loadTranslations()）"],
    ["标记 / 使用方式", "源码写 $localize`:@@id:原文`` 或模板 i18n 属性", "同左（同一套 @@id 标记）", "this.i18n.t('key') / label('key') 按键读取"],
    ["构建产物", "单一 bundle，可切换语言", "n 个独立静态站点（browser/{,zh,de}）", "单一 bundle，可切换 + 可编辑"],
    ["每语言构建次数", "1 次（一份 bundle）", "每种语言各 1 次（localize:true）", "1 次"],
    ["语言切换", "按钮即时切换，无刷新", "不可切换（固定）", "按钮 + 编辑器即时切换，无刷新"],
    ["URL 是否带语言", "不带（/official），内部用 localStorage", "带：/zh/、/de/（base href 指定）", "不带（/custom），内部用 localStorage"],
    ["刷新后语言", "localStorage 记住上次选择", "固定，无所谓刷新", "localStorage 记住上次选择"],
    ["翻译文件格式", "translations.json（经 loadTranslations 注册）", "messages.zh.xlf / messages.de.xlf（<target>）", "translations.json（唯一总表）+ 生成 registry"],
    ["新增语言步骤", "改 JSON 加一列 + 重新构建", "新增 xlf 翻译 + build + 部署新目录", "改 JSON + 重新 codegen"],
    ["包体 / 性能", "含 registry，部分可 tree-shake", "最小、可 tree-shake、无运行时 registry", "含 registry + Monaco 编辑器（较大）"],
    ["类型安全", "无（凭消息 id）", "无", "有（TranslationKey 编译期校验）"],
    ["可在线改文案", "否（需改码 / 改 JSON 重编）", "否（页面已是固定语言）", "是：编辑器 Apply 后即时生效"],
    ["占位符 / ICU", "支持（运行时解析）", "支持（构建时保留占位符）", "支持（复用官方 $localize 机制）"],
    ["Demo 页面", "/official", "/compile-time", "/custom"],
    ["适用场景", "需即时切换、单站点、中小项目", "SEO / 性能、每语言独立发布、内容固定", "文案集中管理、类型安全、快速迭代"],
    ["核心优点", "切换灵活、实现简单", "性能最好、可 tree-shake、无需运行时 registry", "文案集中管理、类型安全、可在线编辑，团队友好"],
    ["核心缺点", "包体与运行时开销略高、加语言要重编", "加语言要重新构建多站点、不能运行时切换", "多一层封装、浏览器里带编辑器偏重"],
]


def style_sheet(ws, headers, rows, widths):
    ws.append(headers)
    for row in rows:
        ws.append(row)
    for col_idx, width in enumerate(widths, start=1):
        ws.column_dimensions[chr(64 + col_idx)].width = width
    for c in ws[1]:
        c.fill = header_fill
        c.font = header_font
        c.alignment = header_align
        c.border = border_all
    for r in range(2, ws.max_row + 1):
        for c in ws[r]:
            c.alignment = wrap
            c.border = border_all
            if r % 2 == 0:
                c.fill = alt_fill
    ws.freeze_panes = "B2"
    ws.auto_filter.ref = ws.dimensions


def build_xlsx():
    wb = Workbook()

    ws1 = wb.active
    ws1.title = "三方案对比"
    style_sheet(ws1, COMPARE_HEADERS, COMPARE_ROWS, widths=[20, 40, 44, 44])

    ws2 = wb.create_sheet("名词小词典")
    style_sheet(ws2, GLOSSARY_HEADERS, GLOSSARY_ROWS, widths=[26, 96])

    ws3 = wb.create_sheet("源码文件地图")
    style_sheet(ws3, FILE_MAP_HEADERS, FILE_MAP_ROWS, widths=[62, 68, 16])

    ws4 = wb.create_sheet("xlf 与 JSON 区别")
    style_sheet(ws4, XLF_JSON_HEADERS, XLF_JSON_ROWS, widths=[18, 48, 48])

    ws5 = wb.create_sheet("extract 格式全景")
    style_sheet(ws5, FORMAT_HEADERS, FORMAT_ROWS, widths=[14, 14, 44, 26, 30, 20])

    ws6 = wb.create_sheet("语言状态与路由")
    style_sheet(ws6, ARCH_HEADERS, ARCH_ROWS, widths=[16, 40, 42])

    ws7 = wb.create_sheet("命令脚本速查")
    style_sheet(ws7, SCRIPTS_HEADERS, SCRIPTS_ROWS, widths=[18, 56, 44, 14])

    ws8 = wb.create_sheet("新增语言 Checklist")
    style_sheet(ws8, CHECKLIST_HEADERS, CHECKLIST_ROWS, widths=[8, 56, 36])

    ws9 = wb.create_sheet("常见问题排查")
    style_sheet(ws9, FAQ_HEADERS, FAQ_ROWS, widths=[30, 40, 56])

    ws10 = wb.create_sheet("标记写法示例")
    style_sheet(ws10, MARK_HEADERS, MARK_ROWS, widths=[14, 60, 46])

    ws11 = wb.create_sheet("xlf 编辑工具")
    style_sheet(ws11, XLF_TOOLS_HEADERS, XLF_TOOLS_ROWS, widths=[20, 16, 18, 34, 22, 42])

    DOC_DIR.mkdir(parents=True, exist_ok=True)
    wb.save(str(XLSX_PATH))
    print(f"[xlsx] saved: {XLSX_PATH}")


if __name__ == "__main__":
    build_docx()
    build_xlsx()
    print("Done.")
